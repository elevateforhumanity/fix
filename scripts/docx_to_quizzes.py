"""
Convert QV5210QuestionBank.docx into per-quiz JSON files.
Expects 4 tables in the DOCX: Core, Type I, Type II, Type III.
Each row: [Quiz label (blank = same quiz), Practice Test #, Answer, Question text]

Usage:
  python3 scripts/docx_to_quizzes.py --docx ./source/QV5210QuestionBank.docx --out ./generated
"""

import argparse
import json
import re
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

from docx import Document
from tqdm import tqdm


def clean(s: str) -> str:
    s = s.replace("\u00a0", " ").replace("\r", " ").replace("\n", " ")
    s = s.replace("\t", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def slugify(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"[\s_-]+", "-", s)
    return s.strip("-") or "untitled"


def split_stem_and_options(raw: str) -> Tuple[str, List[str]]:
    """
    Parse question text like:
      '3) Chlorofluorocarbons ... 1-chlorine... 2-chlorine... 3-carbon...'
    Returns (stem, options[]) where options are in numeric order.
    """
    raw = clean(raw)

    # Drop leading question number "3)" or "12)"
    raw = re.sub(r"^\d+\)\s*", "", raw)

    # Find option markers like "1-" "2-" "3-"
    matches = list(re.finditer(r"(?:^|\s)([1-9])\s*-\s*", raw))
    if not matches:
        return raw, []

    # Stem is everything before first option marker
    first = matches[0]
    stem = raw[: first.start()].strip()

    options: List[str] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(raw)
        opt_text = raw[start:end].strip()
        options.append(opt_text)

    return stem, options


@dataclass
class QuizQuestion:
    id: str
    practice_test: str
    stem: str
    options: List[str]
    correct_option_number: int  # 1-based


@dataclass
class Quiz:
    id: str
    title: str
    section: str
    questions: List[QuizQuestion]


def parse_question_bank(docx_path: Path) -> Dict[str, Quiz]:
    doc = Document(str(docx_path))
    if not doc.tables:
        raise SystemExit("No tables found in DOCX.")

    section_names = ["Core", "Type I", "Type II", "Type III"]
    quizzes: Dict[str, Quiz] = {}

    for ti, table in enumerate(doc.tables):
        section = section_names[ti] if ti < len(section_names) else f"Table {ti+1}"
        current_quiz_label: Optional[str] = None

        for ri, row in enumerate(table.rows):
            cells = [clean(c.text) for c in row.cells]
            if ri == 0:
                continue  # header row

            # Handle tables with 3 or 4 columns
            if len(cells) == 4:
                quiz_cell, practice_test, answer_cell, question_cell = cells
            elif len(cells) == 3:
                quiz_cell, answer_cell, question_cell = cells
                practice_test = ""
            else:
                continue

            if quiz_cell:
                current_quiz_label = quiz_cell

            if not current_quiz_label:
                continue

            if not question_cell:
                continue

            try:
                correct = int(re.sub(r"[^\d]", "", answer_cell))
            except Exception:
                correct = 0

            stem, options = split_stem_and_options(question_cell)

            quiz_num_match = re.search(r"(\d+)", current_quiz_label)
            quiz_num = quiz_num_match.group(1) if quiz_num_match else slugify(current_quiz_label)
            quiz_id = f"{section.lower().replace(' ', '-')}-quiz-{quiz_num}"

            if quiz_id not in quizzes:
                quizzes[quiz_id] = Quiz(
                    id=quiz_id,
                    title=current_quiz_label,
                    section=section,
                    questions=[],
                )

            q_index = len(quizzes[quiz_id].questions) + 1
            qid = f"{quiz_id}-q{q_index:03d}"

            quizzes[quiz_id].questions.append(
                QuizQuestion(
                    id=qid,
                    practice_test=practice_test,
                    stem=stem,
                    options=options,
                    correct_option_number=correct,
                )
            )

    return quizzes


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True, help="Path to QV5210QuestionBank.docx")
    ap.add_argument("--out", required=True, help="Output folder (e.g. ./generated)")
    args = ap.parse_args()

    docx_path = Path(args.docx).resolve()
    out_dir = Path(args.out).resolve()
    quizzes_dir = out_dir / "quizzes"
    quizzes_dir.mkdir(parents=True, exist_ok=True)

    quizzes = parse_question_bank(docx_path)

    manifest: Dict[str, Any] = {
        "source": docx_path.name,
        "quiz_count": len(quizzes),
        "total_questions": sum(len(q.questions) for q in quizzes.values()),
        "quizzes": {},
    }

    for quiz_id, quiz in tqdm(quizzes.items(), desc="Writing quizzes"):
        quiz_obj = asdict(quiz)
        quiz_path = quizzes_dir / f"{quiz_id}.json"
        quiz_path.write_text(json.dumps(quiz_obj, indent=2), encoding="utf-8")

        manifest["quizzes"][quiz_id] = {
            "id": quiz.id,
            "title": quiz.title,
            "section": quiz.section,
            "question_count": len(quiz.questions),
            "path": str(Path("quizzes") / f"{quiz_id}.json").replace("\\", "/"),
        }

    (out_dir / "quizzes.manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    print(f"\nDone. {manifest['total_questions']} questions across {len(quizzes)} quizzes.")
    print(f"- Quiz files: {quizzes_dir}")
    print(f"- Manifest:   {out_dir / 'quizzes.manifest.json'}")


if __name__ == "__main__":
    main()
